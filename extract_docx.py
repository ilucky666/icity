
import docx
import json
import re

def extract_projects(file_path):
    try:
        doc = docx.Document(file_path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                content.append(" | ".join(row_data))
                
        print(json.dumps(content, ensure_ascii=False, indent=2))
        
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    extract_projects(r"F:\LIBRARY\Academic\Urbanrenew\icity\城市更新项目整理.docx")
